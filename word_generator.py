"""Generate Word documents from templates and parsed data."""

import logging
import re
from datetime import date
from pathlib import Path

from docx import Document

from api_client import BureauDB, MockApiClient
from excel_parser import SplitItem

logger = logging.getLogger(__name__)

WAIMAI_COMPANY = "上海京东到家友恒电商信息技术有限公司"
WAIMAI_CREDIT_CODE = "91310110MA1G81DR35"
WAIMAI_ADDRESS = "上海市杨浦区杨树浦路1088号2202单元"
WAIMAI_BUREAU = "上海市杨浦区市场监督管理局"


def _today_str() -> str:
    d = date.today()
    return f"{d.year}年{d.month}月{d.day}日"


def _current_year() -> str:
    return str(date.today().year)


def _format_orders(orders: list[str]) -> str:
    return "、".join(orders)


def _get_bureau(item: SplitItem, db: BureauDB) -> str:
    row = item.row
    cls = row.classification
    mock = MockApiClient()

    if cls == "外卖":
        return WAIMAI_BUREAU

    if cls == "三方公司":
        bureau = db.lookup(row.credit_code)
        if bureau:
            return bureau
        bureau = mock.lookup_by_address(row.address)
        if bureau:
            return bureau
        logger.warning(f"Row {row.row_num}: cannot determine bureau for 三方公司 {row.company}")
        return f"{row.company}对应登记机关"

    if cls == "三方个人":
        bureau = mock.lookup_by_address(row.address)
        if bureau:
            return bureau
        logger.warning(f"Row {row.row_num}: cannot determine bureau for 三方个人 {row.company}")
        return f"{row.address}所属市监局"

    if cls == "自营开票":
        bureau = db.lookup(row.credit_code)
        if bureau:
            return bureau
        bureau = mock.lookup_by_address(row.address)
        if bureau:
            return bureau
        logger.warning(f"Row {row.row_num}: cannot determine bureau for 自营开票 {row.company}")
        return f"{row.company}对应登记机关"

    return ""


def _build_subtitle(task_no: str, prefix4: str) -> str:
    return f"京技管商务消移送〔{_current_year()}〕J{task_no}"


def _build_doc_filename(item: SplitItem) -> str:
    task_no = item.row.task_no_str
    if item.index > 0:
        return f"案件线索移送函J{task_no}-{item.index}"
    return f"案件线索移送函J{task_no}"


def _build_body_text(item: SplitItem) -> str:
    row = item.row
    cls = row.classification
    orders_str = _format_orders(item.order_nos)
    shop = item.shop_name

    if cls == "外卖":
        return (
            f'我单位在举报调查中发现"京东商城（网址: www.jd.com）"'
            f"在销售订单编号{orders_str}的商品时，涉嫌违法的线索。"
            f"经核实，涉案主体为{WAIMAI_COMPANY}，"
            f"注册号/统一社会信用代码为{WAIMAI_CREDIT_CODE}，"
            f"住所在{WAIMAI_ADDRESS}。"
        )

    if cls == "三方公司":
        has_credit = row.credit_code and row.credit_code not in ("暂无", "/", "None", "")
        credit_part = f"，注册号/统一社会信用代码为{row.credit_code}" if has_credit else ""
        return (
            f'我单位在举报调查中发现"京东商城（网址: www.jd.com）"'
            f'上第三方商家"{shop}"在销售订单编号{orders_str}时，'
            f"涉嫌违法的线索。经核实，涉案主体为{row.company}"
            f"{credit_part}，住所在{row.address}。"
        )

    if cls == "三方个人":
        return (
            f'我单位在举报调查中发现"京东商城（网址: www.jd.com）"'
            f'上第三方商家"{shop}"在销售订单编号{orders_str}时，'
            f"涉嫌违法的线索。经核实，涉案主体为{row.company}，"
            f"住所在{row.address}。"
        )

    if cls == "自营开票":
        credit_part = ""
        if row.credit_code and row.credit_code not in ("暂无", "/", "None", ""):
            credit_part = f"，注册号/统一社会信用代码为{row.credit_code}"
        return (
            f'我单位在举报调查中发现"京东商城（网址: www.jd.com）"'
            f"在销售订单编号{orders_str}的商品时，涉嫌违法的线索。"
            f"经核实，涉案主体为{row.company}"
            f"{credit_part}，住所在{row.address}。"
        )

    return ""


def _fill_template(template_path: str, item: SplitItem, db: BureauDB) -> Document:
    doc = Document(template_path)
    row = item.row
    task_no = row.task_no_str
    prefix4 = row.task_no_prefix4

    subtitle = _build_subtitle(task_no, prefix4)
    if item.index > 0:
        subtitle_with_idx = _build_subtitle(f"{task_no}-{item.index}", prefix4)
    else:
        subtitle_with_idx = subtitle

    bureau = _get_bureau(item, db)
    body_text = _build_body_text(item)
    today = _today_str()

    for para in doc.paragraphs:
        text = para.text.strip()

        # Para 2: subtitle
        if text.startswith("京技管商务消移送"):
            for run in para.runs[1:]:
                run.text = ""
            if para.runs:
                para.runs[0].text = f"{subtitle_with_idx}号"
            else:
                para.add_run(f"{subtitle_with_idx}号")
            continue

        # Para 4: notification unit
        if "市场监督管理局" in text and text.endswith("："):
            for run in para.runs[1:]:
                run.text = ""
            if para.runs:
                para.runs[0].text = f"{bureau}：   "
            else:
                para.add_run(f"{bureau}：   ")
            continue

        # Para 5: body text
        if text.startswith("我单位在举报调查中发现"):
            for run in para.runs[1:]:
                run.text = ""
            if para.runs:
                para.runs[0].text = body_text
            else:
                para.add_run(body_text)
            continue

        # Para 12: date
        if re.match(r"\s*\d{4}年\d+月\d+日", text):
            for run in para.runs[1:]:
                run.text = ""
            if para.runs:
                para.runs[0].text = f" {today}"
            else:
                para.add_run(f" {today}")
            continue

    return doc


def generate_word_docs(
    items: list[SplitItem],
    templates: dict[str, str],
    output_dir: str | Path,
    db: BureauDB,
) -> list[dict]:
    output_path = Path(output_dir)
    results = []

    for item in items:
        cls = item.row.classification
        template_path = templates.get(cls)
        if not template_path:
            logger.debug(f"Skipping row {item.row.row_num}: no template for '{cls}'")
            continue

        try:
            doc = _fill_template(template_path, item, db)
            filename = _build_doc_filename(item)
            cls_dir = output_path / cls
            cls_dir.mkdir(parents=True, exist_ok=True)
            out_file = cls_dir / f"{filename}.docx"
            doc.save(str(out_file))

            results.append({
                "row": item.row.row_num,
                "classification": cls,
                "filename": str(out_file),
                "task_no": item.row.task_no_str,
                "shop": item.shop_name,
                "bureau": _get_bureau(item, db),
            })
            logger.info(f"Generated: {out_file}")
        except Exception as e:
            logger.error(f"Failed for row {item.row.row_num}: {e}")
            results.append({
                "row": item.row.row_num,
                "classification": cls,
                "filename": "",
                "error": str(e),
            })

    return results
